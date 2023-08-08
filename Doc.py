import docx # Installed via "pip install python-docx" in cmd Windows
from docx.shared import Pt # Pt size for font

doc = docx.Document('InitialDoc.docx') # Original doc
newdoc = docx.Document()
dash = '–' 

par = newdoc.add_paragraph()
run = par.add_run()  # Создаем объект "Run" для параграфа
run.font.name = 'Comic Sans MS'
run.text = "rjv ujdyf"

#FONT for future==
styleH = doc.styles['Normal']
styleT = doc.styles['Normal']

fontH = styleH.font
fontH.name = 'Arial'
fontH.size = Pt(16)
fontH.bold = True
fontH.italic = False
fontH.underline = None

fontT = styleT.font
fontT.name = 'Arial'
fontT.size = Pt(10)
fontT.bold = False
fontT.italic = False # Maybe None
fontT.underline = None
#==================
#a = ''
for paragraph in doc.paragraphs:
    print(paragraph.text)
    string = str(paragraph.text)
    a = paragraph.text
    if string.count(dash)!= 0 and string != '\n':
        #print("НЕ ЗАГОЛОВОК")
        run.font.size = Pt(10)
        #run.font.name = 'Arial'

    elif string != "\n":
        #print("ЗАГОЛОВОК")
        #print("ooo", string, "ooo")
        run.font.size = Pt(16)
        #run.font.name = 'Arial'

    #run.text = a
    


run.text = "Этот текст будет с другим шрифтом."
#run.text = a
#run.font.size = Pt(16)  # Изменяем размер шрифта для данного объекта "Run"



# Добавляем параграфы с текстом
#newdoc.add_paragraph("Это первый параграф.")
#newdoc.add_paragraph("Это второй параграф.")

# Сохраняем документ в файл
newdoc.save("my_document.docx")